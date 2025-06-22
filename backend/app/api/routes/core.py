from fastapi import APIRouter, Form, UploadFile, File, Request, HTTPException, Depends, status, Response
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from datetime import datetime, timedelta
import hashlib
import json
import tempfile
import shutil
import os
import pickle
import numpy as np
import asyncio
import re
import fitz  # PyMuPDF for PDF processing
import docx
from io import BytesIO
from dotenv import load_dotenv
from app.models.schemas import WelcomeResponse
from app.utils.guardrails import validate_user_input
from app.services.ocr_service import OCRService , split_pdf_to_pages
from app.services.chat_session import chat_session_manager
import tiktoken
from app.db.mongodb import (
    get_user, create_user, update_user_last_login,
    save_document, get_document, save_chat_message,
    get_user_chat_history, save_document_embedding,
    get_document_embeddings, User, Document, ChatMessage,
    DocumentEmbedding, documents_collection, chat_history_collection,
    embeddings_collection, usage_collection, deleted_documents_collection
)
from app.utils.embeddings import get_embedding, cosine_similarity
from fastapi import APIRouter, Depends, HTTPException, status
from typing import List, Dict, Any # Ensure all are imported
from datetime import datetime, timedelta
import traceback
import logging
from openai import AsyncOpenAI

# Correctly import User Pydantic model and get_current_user
# Adjust path if your User model or get_current_user are elsewhere
from app.db.mongodb import User as PydanticUser 
 # Assuming get_current_user is defined here or imported appropriately

from app.services.chat_session import chat_session_manager
# from app.models.schemas import WelcomeResponse # Keep if you have a / route

logger = logging.getLogger(__name__)
router = APIRouter()
# JWT imports
import jwt
import secrets
from passlib.context import CryptContext

# OpenAI imports
from openai import OpenAI

# Load environment variables
load_dotenv()

# Initialize OpenAI client

openai_client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))
# Update model name to a valid one
MODEL_NAME = "gpt-4o-mini"  # or "gpt-4" if you have access

router = APIRouter()

# Authentication setup
SECRET_KEY = os.getenv("SECRET_KEY", "your-very-secret-key")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

# Initialize OCR service
tokenizer = tiktoken.encoding_for_model(MODEL_NAME)
ocr_service = OCRService()

# Constants
MAX_FILE_SIZE_MB = 500  # Increased to 500MB
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024  # Convert to bytes

# Pricing constants (per million tokens or per page)
PRICE_GPT4O_MINI_INPUT_PER_MILLION_TOKENS = 0.15  # $0.15 per million tokens
PRICE_GPT4O_MINI_OUTPUT_PER_MILLION_TOKENS = 0.60  # $0.60 per million tokens
PRICE_OPENAI_EMBEDDING_PER_MILLION_TOKENS = 0.10   # $0.10 per million tokens
PRICE_GEMINI_PRO_OCR_PER_PAGE = 0.0001            # $0.0001 per page (placeholder)

# User models
class UserInDB(User):
    hashed_password: str

class Token(BaseModel):
    access_token: str
    token_type: str

class TokenData(BaseModel):
    username: Optional[str] = None

# Authentication utilities
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

async def authenticate_user(username: str, password: str):
    user = await get_user(username)
    if not user:
        return False
    if not verify_password(password, user.hashed_password):
        return False
    return user

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt



# Utility functions
def calculate_file_hash(file_path: str) -> str:
    """Calculate a hash for a file to use as a unique identifier."""
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

def text_to_chunks(text: str) -> List[str]:
    """
    Splits text into semantic chunks while preserving all content.
    First tries to split by headings, then falls back to paragraphs, then sentences.
    """
    if not text or not text.strip():
        return []
        
    # Clean up the text first
    text = text.replace('\r\n', '\n').replace('\r', '\n').strip()
    
    # Try to split by numbered sections first (e.g., "1. Title" or "## 2. Title")
    section_regex = r'(\n|^)\s*(?:##\s*)?(?:\d{1,2}\.\s+.*?)(?=\n\s*(?:##\s*)?(?:\d{1,2}\.|\Z))'
    sections = re.split(section_regex, text, flags=re.DOTALL)
    
    # If we found sections, clean them up
    if len(sections) > 1:
        chunks = [s.strip() for s in sections if s.strip()]
        # Join section numbers with their content
        chunks = [f"{chunks[i]} {chunks[i+1]}" for i in range(0, len(chunks)-1, 2)]
        if len(chunks) >= 3:  # If we have at least 1.5 sections
            return chunks
    
    # If no sections found or too few, try splitting by double newlines (paragraphs)
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    if len(paragraphs) > 1:
        return paragraphs
    
    # If no paragraphs, try splitting by single newlines
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if lines:
        return lines
    
    # As a last resort, split by sentences
    sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]
    if sentences:
        return sentences
    
    # If all else fails, return the entire text as one chunk
    return [text] if text.strip() else []

def count_tokens(text: str) -> int:
    """Counts the number of tokens in a given text using the model's tokenizer."""
    if not text:
        return 0
    return len(tokenizer.encode(text))
async def generate_and_store_embeddings(document: Document):
    """Generate and store embeddings for document chunks."""
    try:
        # Skip if document already has embeddings for this user
        existing_embeddings = await embeddings_collection.count_documents({
            "document_hash": document.file_hash,
            "user_id": document.user_id
        })
        if existing_embeddings > 0:
            logger.info(f"Document {document.file_hash} already has {existing_embeddings} embeddings for user {document.user_id}")
            return
        
        # Split content into chunks
        chunks = text_to_chunks(document.content)
        logger.info(f"Generated {len(chunks)} chunks for document {document.file_hash}")
        
        # Process chunks in batches to avoid rate limits
        batch_size = 20
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i+batch_size]
            batch_embeddings = []
            
            for j, chunk in enumerate(batch):
                try:
                    # Get embedding for chunk
                    embedding = await get_embedding(chunk)
                    if embedding:
                        # Create embedding document
                        embedding_doc = {
                            "document_hash": document.file_hash,
                            "user_id": document.user_id,
                            "chunk_index": i + j,
                            "chunk_text": chunk,
                            "embedding": embedding,
                            "token_count": count_tokens(chunk),
                            "created_at": datetime.now()
                        }
                        batch_embeddings.append(embedding_doc)
                except Exception as e:
                    logger.error(f"Error generating embedding for chunk {i+j} of document {document.file_hash}: {e}")
            
            # Insert batch of embeddings
            if batch_embeddings:
                await embeddings_collection.insert_many(batch_embeddings)
                logger.info(f"Stored {len(batch_embeddings)} embeddings for document {document.file_hash}")
            
            # Sleep to avoid rate limits
            await asyncio.sleep(1)
        
        # Update document with embedding status
        await documents_collection.update_one(
            {"file_hash": document.file_hash},
            {"$set": {"has_embeddings": True}}
        )
        
        logger.info(f"Successfully generated and stored embeddings for document {document.file_hash}")
    except Exception as e:
        logger.error(f"Error in generate_and_store_embeddings for document {document.file_hash}: {e}", exc_info=True)

# Operation types for usage tracking
class OperationType:
    CHAT = "chat"
    EMBEDDING = "embedding"
    OCR = "ocr"
    TEXT_PROCESSING = "text_processing"
    DOCUMENT_UPLOAD = "document_upload"
    DOCUMENT_DELETE = "document_delete"
    DOCUMENT_REUSE = "document_reuse"

async def log_usage_metrics(
    user_id: str,
    operation: str,
    input_tokens: int = 0,
    output_tokens: int = 0,
    model: str = MODEL_NAME,
    document_count: int = 0,
    document_hashes: List[str] = None,
    page_count: int = 0,
    is_reused: bool = False
):
    """
    Log usage metrics for billing and analytics purposes.
    
    Args:
        user_id: The user ID
        operation: The operation type (use OperationType)
        input_tokens: Number of input tokens
        output_tokens: Number of output tokens
        model: The model used
        document_count: Number of documents involved
        document_hashes: List of document hashes involved
        page_count: Number of pages processed (for OCR)
        is_reused: Whether this is a reused document (no embedding cost)
    """
    try:
        # If document is reused, change operation type for tracking
        if is_reused and operation == OperationType.EMBEDDING:
            operation = OperationType.DOCUMENT_REUSE
            
        usage_log = {
            "user_id": user_id,
            "operation": operation,
            "timestamp": datetime.now(),
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "model": model,
            "document_count": document_count,
            "page_count": page_count,
            "document_hashes": document_hashes or [],
            "is_reused": is_reused
        }
        
        # Calculate estimated costs based on operation type
        if operation == OperationType.CHAT:
            # Chat completion costs
            input_cost = (input_tokens / 1_000_000) * PRICE_GPT4O_MINI_INPUT_PER_MILLION_TOKENS
            output_cost = (output_tokens / 1_000_000) * PRICE_GPT4O_MINI_OUTPUT_PER_MILLION_TOKENS
            usage_log["estimated_cost"] = input_cost + output_cost
            usage_log["cost_breakdown"] = {
                "input_cost": input_cost,
                "output_cost": output_cost,
                "input_tokens": input_tokens,
                "output_tokens": output_tokens
            }
            
        elif operation == OperationType.EMBEDDING:
            # Embedding generation cost (only for new embeddings)
            cost = (input_tokens / 1_000_000) * PRICE_OPENAI_EMBEDDING_PER_MILLION_TOKENS
            usage_log["estimated_cost"] = cost
            usage_log["cost_breakdown"] = {
                "embedding_cost": cost,
                "tokens_processed": input_tokens
            }
            
        elif operation == OperationType.OCR:
            # OCR processing cost (per page)
            cost = page_count * PRICE_GEMINI_PRO_OCR_PER_PAGE
            usage_log["estimated_cost"] = cost
            usage_log["cost_breakdown"] = {
                "ocr_cost": cost,
                "pages_processed": page_count
            }
            
        elif operation == OperationType.DOCUMENT_REUSE:
            # No cost for reusing existing embeddings
            usage_log["estimated_cost"] = 0.0
            usage_log["cost_breakdown"] = {
                "reuse_savings": (input_tokens / 1_000_000) * PRICE_OPENAI_EMBEDDING_PER_MILLION_TOKENS,
                "tokens_saved": input_tokens
            }
            
        # Insert the usage log into the database
        await usage_collection.insert_one(usage_log)
        logger.info(f"Logged {operation} usage for user {user_id}: {input_tokens} in, {output_tokens} out, cost: {usage_log.get('estimated_cost', 0)}")
        
    except Exception as e:
        logger.error(f"Failed to log usage metrics: {e}", exc_info=True)

# Document processing
async def process_document(file_path: str, filename: str, user_id: str) -> Optional[str]:
    """Processes non-PDF text files, now including token counting and usage tracking."""
    try:
        file_hash = calculate_file_hash(file_path)
        
        # Check if document already exists for this user
        existing_doc = await documents_collection.find_one({"file_hash": file_hash, "user_id": user_id})
        if existing_doc:
            # Log document reuse (no cost)
            await log_usage_metrics(
                user_id=user_id,
                operation=OperationType.DOCUMENT_REUSE,
                document_count=1,
                document_hashes=[file_hash],
                input_tokens=existing_doc.get("token_count", 0),
                is_reused=True
            )
            return file_hash

        # Read and process the text document
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text_content = f.read()

        token_count = count_tokens(text_content)
        
        # Create and save the document
        document = Document(
            file_hash=file_hash, 
            filename=filename, 
            user_id=user_id,
            content=text_content, 
            doc_type="text", 
            page_count=1,
            token_count=token_count, 
            embeddings=[]
        )
        await save_document(document)
        
        # Generate embeddings for the document
        await generate_and_store_embeddings(document)
        
        # Log text document processing
        await log_usage_metrics(
            user_id=user_id,
            operation=OperationType.TEXT_PROCESSING,
            document_count=1,
            document_hashes=[file_hash],
            input_tokens=token_count
        )
        
        # Log embedding generation
        await log_usage_metrics(
            user_id=user_id,
            operation=OperationType.EMBEDDING,
            input_tokens=token_count,
            document_hashes=[file_hash]
        )
        
        return file_hash
    except Exception as e:
        logger.error(f"Error processing text document {filename}: {e}", exc_info=True)
        return None

async def process_large_document(file_path: str, filename: str, user_id: str) -> Optional[str]:
    """
    Process document files (PDF or text) including page counting and tokenization.
    
    Args:
        file_path: Path to the file to process
        filename: Original filename
        user_id: ID of the user uploading the file
        
    Returns:
        str: File hash if successful, None otherwise
    """
    try:
        file_hash = calculate_file_hash(file_path)
        
        # Check if document already exists for this user
        existing_doc = await documents_collection.find_one({"file_hash": file_hash, "user_id": user_id})
        if existing_doc:
            # Log document reuse (no cost)
            await log_usage_metrics(
                user_id=user_id,
                operation=OperationType.DOCUMENT_REUSE,
                document_count=existing_doc.get("page_count", 1),
                document_hashes=[file_hash],
                input_tokens=existing_doc.get("token_count", 0),
                page_count=existing_doc.get("page_count", 1),
                is_reused=True
            )
            return file_hash

        # Handle text files
        if filename.lower().endswith(('.txt', '.md', '.csv', '.json', '.yaml', '.yml')):
            logger.info(f"Processing text file: {filename}")
            return await process_document(file_path, filename, user_id)
            
        # Handle PDF files
        elif filename.lower().endswith('.pdf'):
            logger.info(f"Processing PDF file: {filename}")
            return await _process_pdf_document(file_path, filename, user_id, file_hash)
            
        # Handle DOCX files
        elif filename.lower().endswith(('.docx', '.doc')):
            logger.info(f"Processing Word document: {filename}")
            return await _process_docx_document(file_path, filename, user_id, file_hash)
            
        else:
            logger.warning(f"Unsupported file type: {filename}")
            raise ValueError(f"Unsupported file type: {os.path.splitext(filename)[1]}")
            
    except Exception as e:
        logger.error(f"Error processing document {filename}: {e}", exc_info=True)
        return None

async def _process_docx_document(file_path: str, filename: str, user_id: str, file_hash: str) -> Optional[str]:
    """Helper function to process DOCX documents."""
    try:
        # Read the DOCX file
        with open(file_path, "rb") as f:
            docx_bytes = f.read()
        
        # Extract text from DOCX
        doc = docx.Document(BytesIO(docx_bytes))
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        text_content = "\n\n".join(full_text)
        
        # Count pages (estimate: ~500 words per page)
        word_count = len(text_content.split())
        page_count = max(1, (word_count + 499) // 500)  # Round up division
        
        # Count tokens
        token_count = count_tokens(text_content)
        
        # Create document record
        document = Document(
            user_id=user_id,
            filename=filename,
            file_hash=file_hash,
            content=text_content,
            file_size=os.path.getsize(file_path),
            page_count=page_count,
            token_count=token_count,
            is_ocr_processed=False,
            is_digital=True,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )
        
        # Save document to database
        await save_document(document)
        
        # Generate and store embeddings
        await generate_and_store_embeddings(document)
        
        # Log document processing
        await log_usage_metrics(
            user_id=user_id,
            operation=OperationType.DOCUMENT_UPLOAD,
            input_tokens=token_count,
            document_count=1,
            document_hashes=[file_hash],
            page_count=page_count
        )
        
        return file_hash
        
    except Exception as e:
        logger.error(f"Error processing DOCX document {filename}: {e}", exc_info=True)
        return None

async def _process_pdf_document(file_path: str, filename: str, user_id: str, file_hash: str) -> Optional[str]:
    """Helper function to process PDF documents."""
    try:
        # Read the PDF file
        with open(file_path, "rb") as f:
            pdf_bytes = f.read()
        
        # Split PDF into pages for processing
        page_chunks = split_pdf_to_pages(pdf_bytes)
        page_count = len(page_chunks)

        # Try direct text extraction first
        final_text, doc_type = "", "text"
        try:
            with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
                text_from_fitz = "".join(page.get_text() for page in doc)
            
            if len(text_from_fitz.strip()) > 100:
                final_text = text_from_fitz
                doc_type = "text"
                
                # Log text extraction success
                await log_usage_metrics(
                    user_id=user_id,
                    operation=OperationType.TEXT_PROCESSING,
                    document_count=page_count,
                    document_hashes=[file_hash],
                    page_count=page_count
                )
            else:
                # Fall back to OCR if text extraction yields too little content
                final_text, _ = await ocr_service.extract_text_from_pdf(file_path, user_id, file_hash)
                doc_type = "ocr"
                
                # Log OCR processing
                await log_usage_metrics(
                    user_id=user_id,
                    operation=OperationType.OCR,
                    document_count=page_count,
                    document_hashes=[file_hash],
                    page_count=page_count
                )
                
        except Exception as e:
            logger.warning(f"Direct text extraction failed: {e}. Falling back to OCR.")
            final_text, _ = await ocr_service.extract_text_from_pdf(file_path, user_id, file_hash)
            doc_type = "ocr"
            
            # Log OCR processing
            await log_usage_metrics(
                user_id=user_id,
                operation=OperationType.OCR,
                document_count=page_count,
                document_hashes=[file_hash],
                page_count=page_count
            )
        
        if not final_text:
            logger.error(f"Failed to extract text from {filename}")
            return None

        # Count tokens in the extracted text
        token_count = count_tokens(final_text)

        # Create and save the document
        document = Document(
            file_hash=file_hash, 
            filename=filename, 
            user_id=user_id,
            content=final_text, 
            doc_type=doc_type, 
            page_count=page_count,
            token_count=token_count, 
            embeddings=[]
        )
        await save_document(document)
        
        # Generate embeddings for the document
        await generate_and_store_embeddings(document)
        
        # Log embedding generation
        await log_usage_metrics(
            user_id=user_id,
            operation=OperationType.EMBEDDING,
            input_tokens=token_count,
            document_hashes=[file_hash],
            page_count=page_count
        )
        
        return file_hash
        
    except Exception as e:
        logger.error(f"Error processing PDF document {filename}: {e}", exc_info=True)
        return None
# API Endpoints
@router.get("/")
def read_root():
    return WelcomeResponse(
        message="Welcome to the JCS Bot!",
        time=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    )

@router.post("/token", response_model=Token)
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends()):
    user = await authenticate_user(form_data.username, form_data.password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    # Check if the user is an admin
    is_admin = getattr(user, "is_admin", False) or user.username == "admin"
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    # Include the is_admin status in the token data
    token_data = {"sub": user.username, "is_admin": is_admin}
    access_token = create_access_token(
        data=token_data, expires_delta=access_token_expires
    )
    
    await update_user_last_login(user.username)
    
    # Return is_admin in the response body for frontend use
    return {
        "access_token": access_token, 
        "token_type": "bearer",
        "is_admin": is_admin
    }

# UPDATE the get_current_user dependency
async def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        # Extract the is_admin claim from the token payload
        is_admin: bool = payload.get("is_admin", False) 

        if username is None:
            raise credentials_exception
        token_data = TokenData(username=username)
    except jwt.PyJWTError:
        raise credentials_exception
    
    user = await get_user(username=token_data.username)
    if user is None:
        raise credentials_exception
        
    # Attach the is_admin status from the token to the user object
    # This makes it available in dependencies like `admin_required`
    user.is_admin = is_admin
    return user


@router.post("/register")
async def register_user(
    username: str = Form(...),
    password: str = Form(...),
    email: str = Form(...),
    full_name: str = Form(...)
):
    """
    Handles user registration. Expects data as 'multipart/form-data' or 
    'application/x-www-form-urlencoded', which is what HTML forms (and FormData objects) send.
    """
    existing_user = await get_user(username)
    if existing_user:
        raise HTTPException(status_code=400, detail="Username already exists")
    
    # Optional: Check for existing email if it should be unique
    # from app.db.mongodb import get_user_by_email
    # if await get_user_by_email(email):
    #     raise HTTPException(status_code=400, detail="Email already registered")

    hashed_password = get_password_hash(password)
    # Ensure you import User as PydanticUser or use the correct name
    user = PydanticUser(
        username=username,
        email=email,
        full_name=full_name,
        hashed_password=hashed_password,
        disabled=False
    )
    await create_user(user)
    return {"message": "User registered successfully"}

@router.get("/users/me", response_model=User)
async def read_users_me(current_user: User = Depends(get_current_user)):
    return current_user

class FaqRequest(BaseModel):
    prompt: str

# Global context store for maintaining conversation context
conversation_context = {
    'history': [],
    'last_question': None,
    'follow_up': False
}

async def faq_event_generator(prompt: str):
    """
    Enhanced FAQ generator with conversation context and knowledge base integration.
    Maintains conversation history and uses it to improve response quality.
    """
    global conversation_context
    
    # Store current question in context
    conversation_context['last_question'] = prompt
    conversation_context['history'].append(f"User: {prompt}")
    
    # Check if this is a follow-up question
    is_follow_up = conversation_context['follow_up']
    conversation_context['follow_up'] = False  # Reset follow-up flag
    
    try:
        # --- Step 1: Handle simple greetings ---
        greetings = ['hi', 'hii', 'hello', 'hey', 'hai']
        if prompt.lower().strip() in greetings:
            yield "data: " + json.dumps({"chunk": "Hello! How can I help you with our company policies or FAQs today?"}) + "\n\n"
            yield "data: " + json.dumps({"done": True}) + "\n\n"
            return

        admin_user_id = "admin_knowledge_base"
        
        # --- Step 2: Check for knowledge base documents ---
        if await documents_collection.count_documents({"user_id": admin_user_id, "is_knowledge_base": True}) == 0:
            yield "data: " + json.dumps({"chunk": "The knowledge base is currently empty. An administrator needs to upload documents."}) + "\n\n"
            yield "data: " + json.dumps({"done": True}) + "\n\n"
            return

        # --- Step 3: Get all embeddings for the knowledge base ---
        kb_chunks = await embeddings_collection.find({"user_id": admin_user_id}).to_list(length=None)
        if not kb_chunks:
            yield "data: " + json.dumps({"chunk": "No searchable content found. Please ensure documents are processed correctly by an administrator."}) + "\n\n"
            yield "data: " + json.dumps({"done": True}) + "\n\n"
            return

        # --- Step 4: Check for holiday queries first ---
        holiday_keywords = ['holiday', 'holidays', 'time off', 'leave', 'vacation', 'day off']
        is_holiday_query = any(keyword in prompt.lower() for keyword in holiday_keywords)
        
        # Special handling for holiday queries
        if is_holiday_query:
            holiday_text = ""
            for chunk in kb_chunks:
                chunk_text = chunk.get("chunk_text", "").lower()
                if "holidays (2025" in chunk_text or "holiday" in chunk_text:
                    holiday_text = chunk.get("chunk_text", "")
                    break
                    
            if holiday_text:
                # Extract just the holiday section
                holiday_section = holiday_text.split("## 18. Holidays")[1].split("##")[0].strip()
                yield "data: " + json.dumps({"chunk": f"Here are the company holidays for 2025:\n\n{holiday_section}"}) + "\n\n"
                yield "data: " + json.dumps({"done": True}) + "\n\n"
                return
        
        # --- Step 5: Perform similarity search for other queries ---
        prompt_embedding = await get_embedding(prompt)
        if not prompt_embedding:
            raise ValueError("Could not generate embedding for the user's prompt.")
            
        similarities = []
        for chunk in kb_chunks:
            if "embedding" not in chunk or not chunk.get("embedding"):
                continue
                
            chunk_text = chunk.get("chunk_text", "")
            if not chunk_text.strip():
                continue
                
            # Calculate base similarity
            similarity = cosine_similarity(prompt_embedding, chunk["embedding"])
            
            # Convert to lowercase for case-insensitive matching
            chunk_lower = chunk_text.lower()
            prompt_lower = prompt.lower()
            
            # Special handling for policy-related queries
            policy_terms = {
                'working hours': ['timing', 'work schedule', 'office hours', '9:30', '6:30'],
                'leave': ['vacation', 'time off', 'day off', 'holiday', 'casual', 'sick', 'maternity', 'paternity'],
                'remote work': ['wfh', 'work from home', 'remote', 'hybrid', 'work location']
            }
            
            # Check for specific policy terms in the prompt
            for policy, terms in policy_terms.items():
                if any(term in prompt_lower for term in [policy] + terms):
                    # Boost similarity for policy-related chunks
                    if any(term in chunk_lower for term in terms + [policy]):
                        similarity = max(similarity, 0.9)  # Very high boost for direct matches
                        break
            
            # Extract all words from prompt and chunk
            prompt_words = set(re.findall(r'\b\w+\b', prompt_lower))
            chunk_words = set(re.findall(r'\b\w+\b', chunk_lower))
            
            # Calculate word overlap
            common_words = prompt_words.intersection(chunk_words)
            if common_words:
                # Boost based on number of matching words
                word_boost = min(0.5, len(common_words) * 0.1)
                similarity += word_boost
            
            # Boost for any matching content type
            content_boosts = [
                (['ceo', 'chief executive', 'founder', 'leadership'], 0.3),
                (['company', 'about', 'overview', 'nova', 'tech'], 0.2),
                (['product', 'service', 'solution', 'offering'], 0.2),
                (['policy', 'guideline', 'rule', 'process'], 0.2),
                (['value', 'mission', 'vision', 'principle'], 0.2),
                (['event', 'meeting', 'conference', 'holiday'], 0.2),
                (['benefit', 'perk', 'compensation', 'salary'], 0.2)
            ]
            
            for terms, boost in content_boosts:
                if any(term in prompt_lower for term in terms):
                    similarity += boost
                    break
                    
            # Boost for structural elements
            if any(header in chunk_lower for header in ['##', '###', '####']):
                similarity += 0.2
                
            # Boost for direct answers
            if len(chunk_text) < 300 and '?' not in chunk_text:
                similarity += 0.15
                
            # Ensure similarity is within valid range
            similarity = max(0.1, min(1.0, similarity))
                
            # Add to similarities if above minimum threshold
            if similarity > 0.3:  # Lower threshold to include more potential matches
                similarities.append((similarity, chunk_text))
        
        # Sort by similarity score
        similarities.sort(reverse=True, key=lambda x: x[0])
        
        # If no good matches found, try a keyword-based fallback
        if not similarities and len(prompt.split()) < 5:  # For short, direct questions
            keyword_matches = []
            query_terms = set(prompt.lower().split())
            
            for chunk in kb_chunks:
                chunk_text = chunk.get("chunk_text", "")
                if not chunk_text:
                    continue
                    
                chunk_terms = set(chunk_text.lower().split())
                common_terms = query_terms.intersection(chunk_terms)
                if common_terms:
                    # Score based on number of matching terms and chunk length
                    score = len(common_terms) / len(query_terms)
                    if score > 0.3:  # At least 30% of terms match
                        keyword_matches.append((score, chunk_text))
            
            if keyword_matches:
                keyword_matches.sort(reverse=True, key=lambda x: x[0])
                similarities = keyword_matches[:3]  # Take top 3 keyword matches

        # --- Step 5: Get relevant context with enhanced filtering ---
        # Dynamic context window based on query complexity
        is_complex_query = len(prompt.split()) > 5 or '?' in prompt
        top_k = 15 if is_complex_query else 10
        similarity_threshold = 0.30  # Even lower threshold for broader context
        
        # Include previous context for follow-up questions
        context_window = 3  # Number of previous interactions to consider
        previous_context = "\n".join(conversation_context['history'][-context_window:]) if conversation_context['history'] else ""
        
        # Combine current query with context for better embedding
        enhanced_prompt = f"{previous_context}\nCurrent question: {prompt}" if previous_context else prompt
        
        # Include more chunks, even if similarity is lower
        context_chunks = [text for sim, text in similarities[:top_k] if sim > similarity_threshold]
        
        # If we still don't have enough chunks, include more with lower similarity
        if len(context_chunks) < 5 and len(similarities) > len(context_chunks):
            additional_chunks = [text for sim, text in similarities[len(context_chunks):] 
                               if sim > (similarity_threshold * 0.8)][:5]  # Include more chunks with lower similarity
            context_chunks.extend(additional_chunks)
        
        if not context_chunks:
            # If still no chunks, try to include at least the top 3 regardless of score
            context_chunks = [text for _, text in similarities[:3]]
        
        if not context_chunks:
            yield "data: " + json.dumps({"chunk": "I'm sorry, I couldn't find any information related to that topic in the knowledge base."}) + "\n\n"
            yield "data: " + json.dumps({"done": True}) + "\n\n"
            return

        # Add more context about the company to help with general questions
        additional_context = []
        company_info_terms = [
            'ceo', 'company overview', 'about us', 'company name', 'products', 'services',
            'core values', 'values', 'mission', 'vision', 'code of conduct', 'policies',
            'annual events', 'calendar', 'holidays', 'leave policy', 'remote work', 'benefits'
        ]
        
        for chunk in kb_chunks:
            chunk_text = chunk.get("chunk_text", "")
            chunk_lower = chunk_text.lower()
            
            # Check if this chunk contains any company-related information
            if any(term in chunk_lower for term in company_info_terms):
                # Boost relevance of this chunk if it contains section headers
                if any(header in chunk_lower for header in ['##', '###', '####']):
                    if chunk_text not in context_chunks:
                        additional_context.insert(0, chunk_text)  # Add to beginning
                elif chunk_text not in context_chunks:
                    additional_context.append(chunk_text)  # Add to end
        
        # Add up to 3 additional relevant chunks
        context_chunks.extend(additional_context[:3])
        
        # Remove duplicates while preserving order
        seen = set()
        unique_chunks = []
        for chunk in context_chunks:
            if chunk not in seen:
                seen.add(chunk)
                unique_chunks.append(chunk)
        context_chunks = unique_chunks
        
        document_context = "\n\n---\n\n".join(context_chunks)

        # --- Step 6: Prepare the enhanced prompt and call the LLM ---
        
        # Add conversation context to system prompt
        conversation_context_str = "\n".join(conversation_context['history'][-3:])  # Last 3 exchanges
        
        # Enhanced system prompt with conversation awareness
        system_prompt = f"""
        You are NovaAI, the intelligent assistant for NovaTech Inc. You have access to the company's knowledge base and conversation history.
        
        Current Conversation Context:
        {conversation_context_str}
        
        Guidelines:
        1. Be concise but thorough in your responses
        2. If the information is in the knowledge base, provide specific details
        3. For policy questions, include all relevant numbers and conditions
        4. If the question is unclear, ask for clarification
        5. Maintain a helpful and professional tone
        6. When appropriate, suggest related topics the user might find helpful
        
        Knowledge Base Context:
        {document_context}
        
        Current Question: {prompt}
        """
        
        messages = [{"role": "system", "content": system_prompt}]
        
        # Add conversation history if available
        for i in range(0, len(conversation_context['history']), 2):
            if i + 1 < len(conversation_context['history']):
                messages.append({"role": "user", "content": conversation_context['history'][i]})
                messages.append({"role": "assistant", "content": conversation_context['history'][i+1]})
        
        messages.append({"role": "user", "content": prompt})
        policy_answers = {
            'working hours': 'Working hours at NovaTech Inc. are from 9:30 AM to 6:30 PM (Flexible).',
            'leave policy': '''NovaTech Inc. offers the following leave benefits:
            - Casual Leave: 12 Days
            - Earned Leave: 15 Days
            - Sick Leave: 12 Days
            - Maternity Leave: 6 Months
            - Paternity Leave: 15 Days
            Note: Leaves are to be applied via the NovaHR Portal.''',
            'remote work policy': '''Remote Work Policy:
            - Employees can work remotely up to 3 days a week.
            - Permanent remote roles must be approved by HR and the department head.'''
        }
        
        # Check if this is a direct policy question
        prompt_lower = prompt.lower()
        for policy, answer in policy_answers.items():
            if any(term in prompt_lower for term in [policy] + policy.split()):
                yield "data: " + json.dumps({"chunk": answer}) + "\n\n"
                yield "data: " + json.dumps({"done": True}) + "\n\n"
                return
        
        system_prompt = """
        You are a helpful Company FAQ Assistant for NovaTech Inc. Your primary goal is to answer user questions accurately based on the 'Relevant Knowledge Base Context' provided.
        
        Guidelines for responding to common queries:
        
        1. Company Information:
           - Company Name: NovaTech Inc.
           - Industry: Technology Solutions
           - Focus: Innovation, AI, and Enterprise Software
           
        2. Products and Services:
           - AI-powered analytics platforms
           - Cloud computing solutions
           - Enterprise software development
           - IT consulting services
           - Digital transformation solutions
           
        3. Core Values (if not found in context):
           - Innovation: We push boundaries and embrace new ideas
           - Integrity: We uphold the highest ethical standards
           - Excellence: We strive for quality in everything we do
           - Collaboration: We believe in the power of teamwork
           - Customer Focus: We prioritize our clients' success
           
        4. Code of Conduct (if not found in context):
           - Professionalism and respect in all interactions
           - Commitment to diversity and inclusion
           - Protection of company and client confidentiality
           - Compliance with all laws and regulations
           - Zero tolerance for discrimination or harassment
           
        5. Annual Events (if not found in context):
           - Annual General Meeting (Q1)
           - Innovation Summit (Q2)
           - Tech Conference (Q3)
           - Year-End Celebration (Q4)
           
        Response Guidelines:
        - If the information is in the context, provide a detailed answer
        - If context is missing but it's a general query, provide a helpful response based on standard company information
        - Be concise but thorough
        - Include relevant details and examples when possible
        - If unsure, direct to official company resources
        """
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "system", "content": f"Relevant Knowledge Base Context:\n\n{document_context}"},
            {"role": "user", "content": prompt}
        ]
        
        # --- Step 7: Stream the final answer from OpenAI ---
        stream = await openai_client.chat.completions.create(
            model=MODEL_NAME,
            messages=messages,
            stream=True,
            temperature=0.7,  # Increased temperature for more creative/expansive answers
            max_tokens=1000,  # Allow longer responses
            top_p=0.95,  # More diversity in responses
            frequency_penalty=0.2,  # Slightly reduce repetition
            presence_penalty=0.2  # Encourage mentioning of entities
        )
        async for chunk in stream:
            if chunk.choices[0].delta.content:
                yield "data: " + json.dumps({"chunk": chunk.choices[0].delta.content}) + "\n\n"
        
        yield "data: " + json.dumps({"done": True}) + "\n\n"

    except Exception as e:
        logger.error(f"Error in FAQ event generator: {e}", exc_info=True)
        yield "data: " + json.dumps({"error": "An error occurred while generating the response."}) + "\n\n"

@router.post("/faq-chat", tags=["FAQ"])
async def faq_chat(
    request: FaqRequest,
    current_user: PydanticUser = Depends(get_current_user)
):
    prompt = request.prompt
    if not prompt or not prompt.strip():
        raise HTTPException(status_code=400, detail="Prompt cannot be empty.")
        
    return StreamingResponse(faq_event_generator(prompt), media_type="text/event-stream")

@router.post("/chat")
async def chat(
    request: Request,
    prompt: str = Form(...),
    task: Optional[str] = Form(None),
    files: List[UploadFile] = File(None),
    session_id: Optional[str] = Form(None),
    current_user: User = Depends(get_current_user)
):
    try:
        user_id = current_user.username
        original_prompt = prompt
        
        # Enhanced logging for session management
        logger.info(f"Chat endpoint called - User: {user_id}, Session ID: {session_id}, Task: {task}")
        
        session = await chat_session_manager.get_or_create_session(user_id, session_id)
        logger.info(f"Using session {session.session_id} for user {user_id}")
        
        messages = []

        logger.info(f"Task: {task}, Prompt: {original_prompt}, User ID: {user_id}")

        # Track newly processed files separately from context-retrieved documents
        newly_processed_hashes = []
        context_document_hashes = []
        processed_filenames = {}
        
        if files:
            logger.info(f"Processing {len(files)} uploaded files")
            temp_dir = tempfile.mkdtemp()
            try:
                for file in files:
                    if not file.filename: 
                        logger.warning("Skipping file with no filename")
                        continue
                        
                    logger.info(f"Processing file: {file.filename}")
                    file_path = os.path.join(temp_dir, file.filename)
                    with open(file_path, "wb") as f:
                        shutil.copyfileobj(file.file, f)
                    
                    if file.filename.lower().endswith('.pdf'):
                        logger.info(f"Processing PDF file: {file.filename}")
                        file_hash = await process_large_document(file_path, file.filename, user_id)
                    else:
                        logger.info(f"Processing non-PDF file: {file.filename}")
                        file_hash = await process_large_document(file_path, file.filename, user_id)
                        
                    if file_hash:
                        logger.info(f"File processed successfully. Hash: {file_hash}")
                        newly_processed_hashes.append(file_hash)
                        processed_filenames[file_hash] = file.filename
                        if file_hash not in session.active_documents:
                            logger.info(f"Adding document {file_hash} to active documents for session {session.session_id}")
                            session.active_documents.append(file_hash)
                            await session.save_to_db()
                    else:
                        logger.error(f"Failed to process file: {file.filename}")
            finally:
                shutil.rmtree(temp_dir)

        conversation_history = await session.get_conversation_history()
        document_context = ""
        used_documents = []
        
        # Determine if this is a focused task (like summarization)
        is_focused_task = (
            (files and newly_processed_hashes) and
            any(keyword in prompt.lower() for keyword in ["summary", "summarize", "summarise", "explain", "detail"])
        )

        # Fix for Q&A not working on first upload
        # Check if we have active documents but no document context is being retrieved
        if newly_processed_hashes and not is_focused_task:
            logger.info(f"First-time Q&A for newly uploaded documents. Using direct document retrieval.")
            # Get the most recently uploaded document for context
            latest_doc_hash = newly_processed_hashes[-1]
            document_object = await get_document(latest_doc_hash, user_id)
            
            if document_object:
                logger.info(f"Retrieved document {latest_doc_hash} for first-time Q&A")
                # For large documents, we might want to truncate or chunk the content
                document_context = document_object.content[:100000]  # Limit to first 100K chars if very large
                context_document_hashes = [latest_doc_hash]
            else:
                logger.warning(f"Document {latest_doc_hash} not found for first-time Q&A")
        elif is_focused_task:
            doc_hash_to_focus = newly_processed_hashes[-1]
            filename_to_focus = processed_filenames.get(doc_hash_to_focus, "the uploaded document")
            prompt = f"{prompt}: '{filename_to_focus}'" # Make prompt specific
            
            logger.info(f"Focused Task: Getting FULL TEXT of document: {doc_hash_to_focus}")
            document_object = await get_document(doc_hash_to_focus, user_id)
            if document_object:
                # IMPORTANT: Make sure we're getting the full document content
                document_context = document_object.content
                if not document_context or len(document_context) < 100:
                    logger.warning(f"Document content is too short or empty: {len(document_context)} chars")
                    # Try to fetch the document again with full content projection
                    doc = await documents_collection.find_one(
                        {"file_hash": doc_hash_to_focus, "user_id": user_id},
                        projection={"content": 1}
                    )
                    if doc and "content" in doc and doc["content"]:
                        document_context = doc["content"]
                        logger.info(f"Retrieved document content directly: {len(document_context)} chars")
                
                context_document_hashes = [doc_hash_to_focus]
                logger.info(f"Using document content for summarization: {len(document_context)} chars")
            else:
                logger.warning(f"Document {doc_hash_to_focus} not found for focused task")
        elif session.active_documents:
            logger.info("Standard context retrieval: Using similarity search across all active documents.")
            logger.info(f"Active documents for session {session.session_id}: {session.active_documents}")
            
            # Check if embeddings exist for all active documents
            for doc_hash in session.active_documents:
                embedding_count = await embeddings_collection.count_documents({"document_hash": doc_hash})
                logger.info(f"Document {doc_hash} has {embedding_count} embeddings")
                
                # If no embeddings, try to generate them on-the-fly
                if embedding_count == 0:
                    logger.info(f"No embeddings found for document {doc_hash}. Attempting to generate embeddings.")
                    doc = await get_document(doc_hash, user_id)
                    if doc:
                        await generate_and_store_embeddings(doc)
                        logger.info(f"Generated embeddings for document {doc_hash}")
            
            document_context, used_documents = await session.get_document_context_with_sources(prompt)
            context_document_hashes = used_documents
            logger.info(f"Retrieved context from {len(used_documents)} documents using similarity search")

        system_prompt = ""
        if is_focused_task and any(keyword in original_prompt.lower() for keyword in ["summary", "summarize", "summarise"]):
            logger.info("Using multi-record data extraction prompt for summarization.")
            system_prompt = ""
        if is_focused_task and any(keyword in original_prompt.lower() for keyword in ["summary", "summarize", "summarise", "explain", "detail"]):
            logger.info("Using new master prompt for descriptive and structured analysis.")
            # This new "Master Prompt" handles all document types and output styles.
            system_prompt = """
            You are an expert document analyst AI. Your task is to provide a comprehensive analysis of any document provided, following a strict two-part structure. This must be applied to ALL document types, including resumes, invoices, legal affidavits, financial statements, press releases, identification cards, and more.

            **Part 1: Document Description**
            First, begin your response with a concise introductory paragraph. In this paragraph, you MUST identify:
            1. The specific type of document (e.g., "This document is a Curriculum Vitae (CV)...", "This document is a Customer Service Form (CSF)...", "This is an e-PAN Card...").
            2. The document's primary purpose.
            3. The main subject, individual, or company it concerns.

            **Part 2: Detailed Data Extraction**
            After the introductory paragraph, provide a detailed, structured breakdown of the document's contents.
            - Use clear Markdown headings (e.g., `### Personal Information`, `### Academic Qualifications`, `### Order Details`, `### Key Clauses`). Adapt the headings to be relevant to the document's content.
            - Extract all relevant data points under these headings. Be thorough and capture details from every page.
            - For documents containing records of multiple people or items, create a separate, clearly marked section for each one (e.g., `### 1. [Full Name]`, `### 2. [Full Name]`).
            - **CRITICAL:** Do not mix information between different sections or records. Ensure all data is accurately assigned.

            **Example Structure:**

            This document is a [Document Type] for [Main Subject], outlining its [Primary Purpose].

            ### [Relevant Heading 1]
            - **[Data Point A]:** [Extracted Value]
            - **[Data Point B]:** [Extracted Value]

            ### [Relevant Heading 2]
            - **[Data Point C]:** [Extracted Value]
            """
        else:
            logger.info("Using general Q&A prompt with page-awareness.")
            system_prompt = """You are JCS Bot, an advanced enterprise assistant. Your responses should be helpful, informative, and conversational.

When answering questions:
1. If asked for a summary, provide a comprehensive summary of the document content.
2. If the user asks for a PAN number or any personal or sensitive information, and it is present in the provided document context, you MUST extract and return it.
3. Do not refuse to answer if the information is present in the document context, even if it looks like sensitive or private data.
4. Only use the document context for such answers, and do not fabricate information.
5. If the user asks about something in the uploaded documents, always use the document context provided.
6. If the user asks a general question, answer from your own knowledge.
7. If you use information from the document, cite it clearly.
8. Use both document content (if provided) and your general knowledge to give comprehensive answers.
9. Be clear about what information comes from documents vs. your general knowledge.
10. If you're speculating or giving an opinion, make that clear.
11. Your tone should be friendly but professional.
12. Format your responses clearly with good structure.
13. IMPORTANT: When referencing documents, always mention the document name in your response.
14. CRITICAL: When multiple documents are available, ONLY use the context from the most recently uploaded document for summarization or detail explanation tasks, unless explicitly asked about other documents.
15. For each new request, focus ONLY on the document context provided for that specific request.
16. The document text may contain page separators like '--- PAGE BREAK ---'. If the user asks what is on a specific page (e.g., 'what is on page 2?'), use these separators to identify and answer using only the content from that specific page.
"""

        messages = [{"role": "system", "content": system_prompt}]
        
        if document_context:
            document_names = {h: processed_filenames.get(h) for h in newly_processed_hashes}
            for doc_hash in session.active_documents:
                if doc_hash not in document_names:
                    doc = await get_document(doc_hash, user_id)
                    if doc: document_names[doc_hash] = doc.filename
            
            doc_names_str = "\n".join([f"- {h}: {name}" for h, name in document_names.items() if name])
            messages.append({"role": "system", "content": f"Available documents:\n{doc_names_str}\n\nHere is the relevant document context:\n\n{document_context}"})

        # Initialize token counters
        input_tokens = 0
        output_tokens = 0
        
        # Isolate focused tasks from chat history to avoid confusion
        if not is_focused_task and conversation_history:
            logger.info("Adding conversation history to the prompt.")
            messages.append({"role": "system", "content": f"Here is the recent chat history:\n\n{conversation_history}"})

        full_prompt_for_api = " ".join([m["content"] for m in messages]) + prompt
        input_tokens = count_tokens(full_prompt_for_api)

        # Use the potentially modified prompt for the LLM call
        messages.append({"role": "user", "content": prompt})
        
        async def generate_response():
            full_response_text = ""
            try:
                stream = await openai_client.chat.completions.create(model=MODEL_NAME, messages=messages, stream=True)
                async for chunk in stream:
                    content = chunk.choices[0].delta.content
                    if content:
                        full_response_text += content
                        yield f"data: {json.dumps({'chunk': content})}\n\n"
            except Exception as e:
                logger.error(f"OpenAI error: {e}", exc_info=True)
                full_response_text = f"Error: {e}"
                yield f"data: {json.dumps({'error': str(e)})}\n\n"

            output_tokens = count_tokens(full_response_text)
            
            all_referenced_hashes = list(set(newly_processed_hashes + context_document_hashes))
            logger.info(f"Final referenced documents for this message: {all_referenced_hashes}")
            
            # Save the message to the session
            await session.add_message(
                original_prompt, full_response_text, all_referenced_hashes, 
                {}, input_tokens=input_tokens, output_tokens=output_tokens
            )
            
            # Log chat usage metrics
            try:
                await log_usage_metrics(
                    user_id=user_id,
                    operation=OperationType.CHAT,
                    input_tokens=input_tokens,
                    output_tokens=output_tokens,
                    document_count=len(all_referenced_hashes),
                    document_hashes=all_referenced_hashes
                )
                logger.info(f"Logged chat usage: {input_tokens} input tokens, {output_tokens} output tokens")
            except Exception as e:
                logger.error(f"Failed to log chat usage metrics: {e}", exc_info=True)
            
            # Send session info before completion
            session_data = {
                'session_id': session.session_id,
                'active_documents': session.active_documents
            }
            yield f"data: {json.dumps(session_data)}\n\n"
            
            # Signal completion
            yield f"data: {json.dumps({'done': True})}\n\n"

        return StreamingResponse(generate_response(), media_type="text/event-stream")

    except Exception as e:
        logger.error(f"Error in chat endpoint: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to process request.")


@router.get("/health")
async def health_check():
    """Check if the API is running."""
    return {
        "status": "healthy",
        "time": datetime.now().isoformat()
    }

@router.get("/session/{session_id}")
async def get_session(session_id: str, current_user: User = Depends(get_current_user)):
    """Get chat session details."""
    try:
        session = await chat_session_manager.get_or_create_session(current_user.username, session_id)
        return {
            "session_id": session.session_id,
            "active_documents": session.active_documents,
            "last_activity": session.last_activity,
            "chat_history": session.chat_history
        }
    except Exception as e:
        print(f"Error getting session: {e}")
        raise HTTPException(status_code=500, detail=str(e))



@router.get("/documents", response_model=List[Document])
async def get_user_documents(current_user: User = Depends(get_current_user)):
    """Get all documents for the current user."""
    try:
        # Get documents for the current user
        cursor = documents_collection.find({"user_id": current_user.username})
        documents = await cursor.to_list(length=None)
        
        # Convert to Document models
        return [Document(**doc) for doc in documents] if documents else []
    except Exception as e:
        print(f"Error fetching documents: {str(e)}")  # Add logging
        raise HTTPException(
            status_code=500,
            detail="Failed to fetch documents. Please try again later."
        )

@router.delete("/documents/{file_hash}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(file_hash: str, current_user: User = Depends(get_current_user)):
    """
    Delete a document and its associated embeddings.
    Also tracks the deletion for billing purposes.
    """
    # Get the document first to check ownership and save details
    document = await documents_collection.find_one({
        "file_hash": file_hash,
        "user_id": current_user.username
    })
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found or access denied"
        )
    
    # Save document details before deletion for billing purposes
    deleted_doc = {
        "user_id": current_user.username,
        "file_hash": file_hash,
        "filename": document.get("filename", ""),
        "doc_type": document.get("doc_type", "ocr"),
        "page_count": document.get("page_count", 0),
        "token_count": document.get("token_count", 0),
        "deleted_at": datetime.utcnow()
    }
    
    # Delete document and its embeddings
    await asyncio.gather(
        documents_collection.delete_one({"file_hash": file_hash, "user_id": current_user.username}),
        embeddings_collection.delete_many({"document_hash": file_hash, "user_id": current_user.username}),
        deleted_documents_collection.insert_one(deleted_doc)
    )
    
    return Response(status_code=status.HTTP_204_NO_CONTENT)

@router.delete("/session/{session_id}")
async def delete_chat_session_route(session_id: str, current_user: PydanticUser = Depends(get_current_user)): # Renamed for clarity
    """End and delete a specific chat session for the current user."""
    if not current_user or not current_user.username:
        logger.error("In /session/{session_id} DELETE: current_user or current_user.username is invalid.")
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid user.")

    logger.info(f"User {current_user.username} attempting to delete session: {session_id}")
    try:
        # --- MODIFIED HERE: Pass current_user.username ---
        await chat_session_manager.end_session(session_id, current_user.username)
        logger.info(f"Session {session_id} processed for deletion for user {current_user.username}.")
        return {"message": "Session deleted successfully"}
    except Exception as e:
        error_traceback = traceback.format_exc()
        logger.critical(
            f"CRITICAL ERROR in DELETE /session/{session_id} for user {current_user.username}. Error: {str(e)}\nTRACEBACK:\n{error_traceback}"
        )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An internal server error occurred while deleting the session."
        )

@router.get("/sessions", response_model=Dict[str, List[Dict[str, Any]]])
async def get_user_sessions_list_route(current_user: PydanticUser = Depends(get_current_user)):
    if not current_user or not current_user.username:
        logger.error("In /sessions endpoint: current_user or current_user.username is invalid.")
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid user information for fetching sessions."
        )
    
    logger.info(f"Endpoint /sessions called by user: {current_user.username}")
    try:
        sessions_data = await chat_session_manager.get_user_sessions(current_user.username)
        logger.info(f"ChatSessionManager returned {len(sessions_data)} sessions for user {current_user.username} to /sessions endpoint.")
        return {"sessions": sessions_data}

    except HTTPException as he:
        logger.error(f"HTTPException in /sessions for user {current_user.username}: {he.detail}", exc_info=False)
        raise he
    except Exception as e:
        error_traceback = traceback.format_exc()
        logger.critical(
            f"CRITICAL UNHANDLED ERROR in /sessions endpoint for user {current_user.username}. Error: {str(e)}\nTRACEBACK:\n{error_traceback}"
        )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, 
            detail="An internal server error occurred while fetching your chat sessions. Please check server logs for details."
        )
    